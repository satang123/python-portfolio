from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)


document = Document()

#profile picture
document.add_picture('satex.jpeg', width=Inches(2.0))

#name phone number and email details
name = input('What is your name? ')
speak('Hello ' + name + 'how are you today?')
phone_number = input('What is your phone number? ')
email = input('What is your email address? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email
)

#about me
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself ')
)

#work experience
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter companyh ')
from_date = input('from Date ')
to_date = input('to Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company
)
p.add_run(experience_details)

document.save('cv.docx')